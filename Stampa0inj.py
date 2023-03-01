import glob
import os
import time
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx2pdf import convert

from Unifilare import *


def documento_0inj(dict_all, serie, fase, sonda):
    document = Document()
    header = document.sections[0].header
    htable = header.add_table(1, 2, Inches(6))
    htab_cells = htable.rows[0].cells
    ht0 = htab_cells[0].add_paragraph()
    kh = ht0.add_run()
    kh.add_picture('./img/CartaIntestata/header.png', width=Inches(6))
    footer = document.sections[0].footer
    ftable = footer.add_table(1, 2, Inches(6))
    ftab_cells = ftable.rows[0].cells
    ft0 = ftab_cells[0].add_paragraph()
    fh = ft0.add_run()
    fh.add_picture('./img/CartaIntestata/footer.png', width=Inches(6))

    document.add_heading('Configurazione Zero-Injection per inverter ' + str(serie), 0)
    if sonda == 'METER':
        pin = dict_all[fase][serie]["0-INJ"]["Pin Meter"]
        pin_sonda = ''
        pin_inv = ''
        for i in pin:
            if "Meter" in i:
                pin_sonda += i + ': ' + str(dict_all[fase][serie]["0-INJ"]["Pin Meter"][i]) + '\n'
            elif "Inverter" in i:
                pin_inv += i + ': ' + str(dict_all[fase][serie]["0-INJ"]["Pin Meter"][i]) + '\n'

        if '1PH' in fase:
            document.add_heading('Inverter + ZSM-METER-DDSU', level=1)
        else:
            document.add_heading('Inverter + ZSM-METER-DTSU', level=1)
    elif sonda in ['TA', 'CT']:
        document.add_heading('Inverter con Sensori di Corrente', level=1)
        pin = dict_all[fase][serie]["0-INJ"]["Pin TA"]
        pin_sonda = ''
        pin_inv = ''
        for i in pin:
            if "TA" in i:
                pin_sonda += i + ': ' + str(dict_all[fase][serie]["0-INJ"]["Pin TA"][i]) + '\n'
            elif "Inverter" in i:
                pin_inv += i + ': ' + str(dict_all[fase][serie]["0-INJ"]["Pin TA"][i]) + '\n'
    elif sonda == 'ENERCLICK':
        pin = dict_all[fase][serie]["0-INJ"]["Pin ENERCLICK"]
        pin_com_rs485 = ''
        pin_com_inv = ''
        for i in pin:
            if "Meter" in i:
                pin_com_rs485 += i + ': ' + str(dict_all[fase][serie]["0-INJ"]["Pin ENERCLICK"][i]) + '\n'
            elif "Inverter" in i:
                pin_com_inv += i + ': ' + str(dict_all[fase][serie]["0-INJ"]["Pin ENERCLICK"][i]) + '\n'
        pin = dict_all[fase][serie]["0-INJ"]["Pin Meter"]
        pin_sonda = ''
        pin_inv = ''
        for i in pin:
            if "Meter" in i:
                pin_sonda += i + ': ' + str(dict_all[fase][serie]["0-INJ"]["Pin Meter"][i]) + '\n'
            elif "Inverter" in i:
                pin_inv += i + ': ' + str(dict_all[fase][serie]["0-INJ"]["Pin Meter"][i]) + '\n'
    document.add_paragraph('\n')

    if sonda != 'ENERCLICK':
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'DEVICE'
        hdr_cells[1].text = 'PIN'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Inverter'
        row_cells[1].text = pin_inv
        row_cells = table.add_row().cells
        row_cells[0].text = str(sonda)
        row_cells[1].text = pin_sonda
    else:
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'DEVICE'
        hdr_cells[1].text = 'PIN'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Inverter'
        row_cells[1].text = pin_inv
        row_cells = table.add_row().cells
        row_cells[0].text = 'METER'
        row_cells[1].text = pin_sonda

    if sonda == 'METER':
        if "HYD" in serie and "3" in fase:
            if '5 e 6' in table.cell(1, 1).text:
                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture('./0inj/misc/COM/hyd-3ph-hp-meter1.png', width=Inches(2.55))
                run_2 = paragraph.add_run()
                run_2.add_picture('./0inj/misc/COM/hyd-3ph-hp-meter2.png', width=Inches(2.55))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "HYD" in serie and "1" in fase:
            if 'ZCS-1PH-HYD-3000_6000-ZSS' in serie:
                document.add_picture('./0inj/misc/COM/hyd-1ph-zss-meter.png', width=Inches(2.55))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if 'ZCS-1PH-HYD-3000_6000-ZSS-HP' in serie:
                document.add_picture('./0inj/misc/COM/hyd-1ph-hp-meter.png', width=Inches(2.55))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif serie in ['ZCS-3PH-3.3_12KTL-V3', 'ZCS-3PH-15000_24000TL-V3', 'ZCS-3PH-80_110KTL-LV',
                       'ZCS-3PH-100_136KTL-HV',
                       'ZCS-3PH-3.3_12KTL-V1', 'ZCS-3PH-50000_60000TL-V1', 'ZCS-3PH-20000_33000TL-V2',
                       'ZCS-3PH-10_15KTL-V2', 'ZCS-3PH-25_50KTL-V3'
                       ]:
            if '5 e 6' in table.cell(1, 1).text:
                document.add_picture('./0inj/misc/COM/v3.png', width=Inches(2.55))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif sonda == 'TA':
        if "HYD" in serie and "3" in fase:
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.add_picture('./0inj/misc/COM/hyd-3ph-ta.png', width=Inches(2.55))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "HYD" in serie and "1" in fase:
            if 'ZCS-1PH-HYD-3000_6000-ZSS' in serie:
                document.add_picture('./0inj/misc/COM/hyd-1ph-zss-ta.png', width=Inches(2.55))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if 'ZCS-1PH-HYD-3000_6000-ZSS-HP' in serie:
                document.add_picture('./0inj/misc/COM/hyd-1ph-hp-ta.png', width=Inches(2.55))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif serie in ['ZCS-3PH-3.3_12KTL-V3', 'ZCS-3PH-15000_24000TL-V3', 'ZCS-3PH-80_110KTL-LV',
                       'ZCS-3PH-100_136KTL-HV',
                       'ZCS-3PH-3.3_12KTL-V1', 'ZCS-3PH-50000_60000TL-V1', 'ZCS-3PH-20000_33000TL-V2',
                       'ZCS-3PH-10_15KTL-V2', 'ZCS-3PH-25_50KTL-V3'
                       ]:
            if '5 e 6' in table.cell(1, 1).text:
                document.add_picture('./0inj/misc/COM/v3.png', width=Inches(2.55))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif sonda == 'ENERCLICK':
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture('./0inj/misc/COM/ComBox_meter.png', width=Inches(2.55))
        # run_2 = paragraph.add_run()
        # run_2.add_picture('./0inj/misc/COM/hyd-3ph-hp-meter2.png', width=Inches(2.55))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    unifilare(serie, fase, sonda)
    time.sleep(0.5)
    document.add_heading('Unifilare', level=1)
    document.add_paragraph('\n')
    document.add_picture('./0inj/img/Schema_' + serie + '_' + sonda + '.png', width=Inches(4.75))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sonda == 'METER' or sonda == 'ENERCLICK':
        document.add_page_break()
        if sonda == 'ENERCLICK':
            document.add_heading('Guida alla configurazione del ComBox', level=1)
            document.add_paragraph(
                'Il modulo ComBox, oltre a poter controllare la Potenza prodotta dagli inverter, può '
                'effettuare il monitoraggio dei consumi dell’impianto. L\'installazione può essere '
                'eseguita utilizzando i dongle ETH (ZSM-ETH-USB, uno per inverter) oppure utilizzando '
                'la porta RS485 dell’Inverter. \nLa configurazione mediante dongle ETH, permette un'
                ' monitoraggio più accurato dei singoli Inverter ed è fortemente consigliata.')
            document.add_heading('Installazione e configurazione impianto con Combox-Comunicazione ETH', level=2)
            document.add_heading('L\'impianto deve essere cablato come lo schema seguente\n')
            document.add_picture('./0inj/misc/Enerclick/ENERCLICK_ETH.png', width=Inches(4.75))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph('Gli inverter devono essere configurati con IP statico', style='List Bullet')
            document.add_paragraph('Inserire 2 resistenze di terminazione da 120 Ohm sul pin 24-25 del meter e sulla '
                                   'coppia A1-B1 del ComBox', style='List Bullet')
            document.add_paragraph('Verificare che la porta 80 del router sia aperta', style='List Bullet')

            document.add_heading('Installazione e configurazione impianto con Combox-Comunicazione RS485', level=2)
            document.add_paragraph('L\'impianto deve essere cablato come lo schema seguente\n')
            document.add_picture('./0inj/misc/Enerclick/ENERCLICK_485.png', width=Inches(4.75))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph('Inserire 4 resistente di terminazione da 120Ω sui pin 24-25 del meter, sulla'
                                   ' coppia di pin A1-B1 e A2-B2 del ComBox e sui pin della porta RS485 dell\'inverter '
                                   'qualora la lunghezza dei cavi ecceda i 20m', style='List Bullet')
            document.add_paragraph('Verificare che la porta 80 del router sia aperta', style='List Bullet')

            document.add_heading('Installazione e configurazione ComBox', level=2)
            document.add_paragraph('Lo strumento verrà venduto con un generatore 5V a barra DIN. Il ComBox va collegato '
                                 'come in figura\n')
            document.add_picture('./0inj/misc/Enerclick/Alim_ENERCLICK.png', width=Inches(4.75))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_page_break()

        document.add_heading('Guida alla configurazione del Meter', level=1)
        document.add_paragraph('\nCablare il meter come risportato in figura:\n')
        if "3" in fase:
            document.add_picture('./0inj/misc/meter/Meter+CT-3ph.png', width=Inches(4.75))
        elif "1" in fase:
            document.add_picture('./0inj/misc/meter/Meter+CT-ph.png', width=Inches(4.75))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('\nCablare il meter all\'inverter come riportato nella tabella sopra riportata.'
                               '\nPer configurare il meter DTSU è necessario fare le seguenti operazioni:')
        document.add_paragraph(
            'Premere SET sul meter. ', style='List Bullet'
        )
        document.add_paragraph(
            'Quando comparirà la scritta CODE, premere SET una seconda volta. ', style='List Bullet'
        )
        document.add_paragraph(
            'Scrivere 701 usando il tasto SET per spostare il cursore verso sinistra e “→” per aumentare di 1 la cifra'
            ' selezionata. In caso di errore premere ESC e poi di nuovo SET per reimpostare il codice voluto.',
            style='List Bullet'
        )
        document.add_picture('./0inj/misc/meter/PortaCHINT.png', width=Inches(4.75))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(
            'Confermare il numero premendo SET più volte finché non compare la scritta CT.',
            style='List Bullet'
        )
        document.add_paragraph(
            'Premere SET per configurare il rapporto primario/secondario del CT.',
            style='List Bullet'
        )
        document.add_paragraph(
            'Impostare il valore del rapporto primario/secondario nella stessa maniera con cui è stato configurato '
            'il valore precedente. Per le nostre sonde, il rapporto è 40.',
            style='List Bullet'
        )
        document.add_paragraph(
            'Premere ESC per confermare e “→” per scorrere fino all’impostazione ADDR.',
            style='List Bullet'
        )
        document.add_picture('./0inj/misc/meter/ctCHINT.png', width=Inches(4.75))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(
            'Premere SET per configurare l\'indirizzo del meter.',
            style='List Bullet'
        )
        document.add_paragraph(
            'Se il meter in oggetto è il meter di scambio, lasciare “01” come indirizzo.\nSe il meter è un meter non di'
            ' scambio, selezionare il numero successivo. In questa modo, l\'inverter assegnerà come potenze relative '
            'alla produzione i dati inviati dal meter. Possono essere collegati massimo 3 meter non di scambio, '
            'ovvero fino all\'indirizzo "04" (da "02 a "04").',
            style='List Bullet'
        )
        document.add_paragraph(
            'Premere ESC per confermare e ESC di nuovo per uscire dal menù.',
            style='List Bullet'
        )
        document.add_picture('./0inj/misc/meter/AdressCHINT.png', width=Inches(4.75))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if sonda == 'ENERCLICK':
        document.add_page_break()
        document.add_heading('Guida alla configurazione del ComBox', level=1)
        document.add_heading('Configurazione del Meter', level=2)
        document.add_paragraph('\n')
        file = open('./0inj/misc/Enerclick/1_Combox_meter/istruzioni.txt', 'r')
        lines = file.readlines()
        j = 0
        for k in glob.glob('./0inj/misc/Enerclick/1_Combox_meter/*.png'):
            document.add_picture(k, width=Inches(2.75))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(lines[j])
            j += 1

        document.add_heading('Configurazione del Combox - RS485', level=2)
        document.add_paragraph('\n')
        file = open('./0inj/misc/Enerclick/2_Combox_Inverter_RS485/istruzioni.txt', 'r')
        lines = file.readlines()
        j = 0
        for k in glob.glob('./0inj/misc/Enerclick/2_Combox_Inverter_RS485/*.png'):
            document.add_picture(k, width=Inches(2.75))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(lines[j])
            j += 1

        document.add_heading('Configurazione del Combox - TCP', level=2)
        document.add_paragraph('\n')
        file = open('./0inj/misc/Enerclick/3_Combox_Inverter_TCP_IP/istruzioni.txt', 'r')
        lines = file.readlines()
        j = 0
        for k in glob.glob('./0inj/misc/Enerclick/3_Combox_Inverter_TCP_IP/*.png'):
            document.add_picture(k, width=Inches(2.75))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(lines[j])
            j += 1

        document.add_heading('Configurazione del Combox - Finale', level=2)
        document.add_paragraph('\n')
        file = open('./0inj/misc/Enerclick/4_Combox_Configurazione finale/istruzioni.txt', 'r')
        lines = file.readlines()
        j = 0
        for k in glob.glob('./0inj/misc/Enerclick/4_Combox_Configurazione finale/*.png'):
            document.add_picture(k, width=Inches(4.75))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(lines[j])
            j += 1

    elif sonda == 'TA':
        document.add_page_break()
        document.add_heading('Guida alla configurazione del TA', level=1)

        document.add_paragraph('\nVerificare che la versione del Firmware dell\'inverter sia pari o superiore alla '
                               '70008. Per verificare la versione, tramite il display dell\'inverter:\n')
        document.add_paragraph('Premere il primo tasto a sinistra', style='List Bullet')
        document.add_paragraph('Andare su "Info Sistema" e premere la freccia ->', style='List Bullet')
        document.add_paragraph('Andare su "Info Inverter" e premere la freccia ->', style='List Bullet')
        document.add_paragraph('Scorrere fino alla voce "Codice Servizio" e premere la freccia ->', style='List Bullet')
        document.add_paragraph('Digitare la PSW 0715', style='List Bullet')
        document.add_paragraph('\nSe la versione del Firmeare è inferiore alla 70008, aggiornare l\'inverter. '
                               'Se la versione del Firmware dell\'inverter è pari o superiore alla '
                               '70008:\n')
        document.add_paragraph('Ruotare su OFF l\'interruttore sotto l\'inverter per la parte fotovoltaica',
                               style='List Bullet')
        document.add_paragraph('Tenere su ON le Batterie e lato AC', style='List Bullet')
        document.add_paragraph('Sganciare tutti i carichi dell\'utenza', style='List Bullet')
        document.add_paragraph('Sganciare tutte le sonde dalle fasi se già installate', style='List Bullet')
        document.add_paragraph('Impostare l\'inverter in scarica forzata a 3kW', style='List Bullet')
        document.add_paragraph('Avendo un sistematrifase perfettamente bilanciato, dall’inverter stesso, la potenza '
                               'sarà suddivisa sulle fasi equamente', style='List Bullet')
        document.add_paragraph('Premere dalla schermata iniziale dell\'inverter due volte la freccia in giu\'',
                               style='List Bullet')
        document.add_paragraph('Qui è possible visualizza il PF di tutte le fasi su N/A in quanto i CT non sono '
                               'ancora agganciati alle fasi.', style='List Bullet')
        document.add_paragraph('Installare un CT alla volta intorno alla singola fase controllando sull\'inverter il '
                               'PF. il PF deve assumere valore simile a 1 se si è collegati alla fase giusta, in caso '
                               'contrario collegarsi al un altra fase. Ripere questa manovra su tutti e 3 i CT fino ad '
                               'ottenere tutti i CT con valori simili a 1 e dicitura EXPORT.', style='List Bullet')
        document.add_paragraph('Riarmare l\'inverter partendo dallo switch fotovoltaico, le batterie e poi il generale '
                               'lato AC.', style='List Bullet')
        # document.add_paragraph('Impostare in maniera automatica le batterie come da manuale.', style='List Bullet')

    document.add_page_break()
    if sonda != 'ENERCLICK':
        document.add_heading('Guida alla configurazione dell\'inverter', level=1)
        if "HYD" in serie and "3" in fase:
            document.add_paragraph('\n')
            file = open('./0inj/misc/hyd 3ph/istruzioni.txt', 'r')
            lines = file.readlines()
            j = 0
            for i in glob.glob('./0inj/misc/hyd 3ph/*.png'):
                document.add_picture(i, width=Inches(4.75))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph(lines[j])
                j += 1
        elif "HYD" in serie and "1" in fase:
            document.add_paragraph('\n')
            file = open('./0inj/misc/hyd 1ph/istruzioni.txt', 'r')
            lines = file.readlines()
            j = 0
            for i in glob.glob('./0inj/misc/hyd 1ph/*.png'):
                document.add_picture(i, width=Inches(4.75))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph(lines[j])
                j += 1
        elif serie in ['ZCS-3PH-3.3_12KTL-V3', 'ZCS-3PH-15000_24000TL-V3', 'ZCS-3PH-80_110KTL-LV', 'ZCS-3PH-100_136KTL-HV',
                       'ZCS-3PH-3.3_12KTL-V1', 'ZCS-3PH-50000_60000TL-V1', 'ZCS-3PH-20000_33000TL-V2',
                       'ZCS-3PH-10_15KTL-V2',
                       'ZCS-3PH-25_50KTL-V3',
                       'ZCS-1PH-3000_6000TLM-V3'
                       ]:

            document.add_paragraph('\n')
            file = open('./0inj/misc/v3/istruzioni.txt', 'r')
            lines = file.readlines()
            j = 0
            for i in glob.glob('./0inj/misc/v3/*.jpg'):
                document.add_picture(i, width=Inches(4.75))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph(lines[j])
                j += 1
        elif serie in ['ZCS-SOFAR-10000_20000TL', 'ZCS-1PH-1100_3300TL-V3', 'ZCS-1PH-1100_3300TL-V1',
                       'ZCS-1PH-3000_6000TLM-V2']:
            document.add_paragraph('\nPremere in maniera prolungata (3 secondi) il pulsante quando ci si trova nella '
                                   'schermata dell’interfaccia principale per accedere al menù principale\nSelezionare '
                                   '“18. Contr P(rete)” e premere il pulsante in maniera prolungata '
                                   'per accedere al relativo menù, da cui sarà possibile abilitare la funzione di '
                                   '“Reflux Power”, ovvero la possibilità di impostare la massima potenza immessa in '
                                   'rete. A display apparirà l’indicazione “Immettere PWD!”, quindi digitare la'
                                   ' password “0001” premendo brevemente il pulsante per incrementare la cifra '
                                   'selezionata e premendo in maniera prolungata il pulsante per passare a quella '
                                   'seguente e confermare. Se dovesse comparire a display l’indicazione “Errore, '
                                   'riprova!”, premere nuovamente il tasto e digitare nuovamente la password. Quando la '
                                   'password digitata risulterà corretta, sarà possibile accedere al menù. A questo '
                                   'punto, premendo brevemente il pulsante sarà possibile scegliere l’opzione '
                                   '“1.Abilita” o “2.Disabilita” e selezionarla premendo il pulsante in maniera '
                                   'prolungata. Se viene selezionata l’opzione “1.Abilita”, premendo brevemente il'
                                   ' pulsante sarà possibile selezionare il valore di potenza (espresso in kW, fino al'
                                   ' secondo decimale) di potenza massima che l’inverter andrà a immettere in rete; in'
                                   ' tal modo l’inverter potrà immettere nella rete elettrica nazionale un quantità di'
                                   ' potenza massima compresa tra 0 kW e la potenza nominale dell’inverter, sulla base'
                                   ' della radiazione solare disponibile e dei consumi domestici. A display sarà '
                                   'visualizzata l’indicazione “OK” se l’impostazione sarà andata a buon fine; in caso '
                                   'contrario, sarà visualizzata l’indicazione “Errore”.')
    document.save('./0inj/' + serie + '_' + sonda + '.docx')
    time.sleep(1)
    convert('./0inj/' + serie + '_' + sonda + '.docx', './0inj/' + serie + '_' + sonda + '.pdf')
    time.sleep(1)
    os.remove('./0inj/' + serie + '_' + sonda + '.docx')
